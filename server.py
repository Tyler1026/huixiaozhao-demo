# deploy v20260826-1535
#!/usr/bin/env python3
"""
慧小招 KB Chat Server — localhost:5050
- GET  /            → 返回 HTML demo 文件
- GET  /health      → 健康检查
- POST /api/kb-chat → RAG 问答，流式调用 DeepSeek API
"""
import json, os, urllib.request, urllib.error, urllib.parse, time, datetime, hashlib, base64, io
from http.server import HTTPServer, ThreadingHTTPServer, BaseHTTPRequestHandler
try:
    import psycopg2
    _PG_AVAIL = True
except ImportError:
    _PG_AVAIL = False

# 文档文本提取：PDF 用 pdfplumber，Word(.docx) 用 python-docx；缺库时优雅降级
try:
    import pdfplumber
    _PDF_AVAIL = True
except ImportError:
    _PDF_AVAIL = False
try:
    import docx as _docx
    _DOCX_AVAIL = True
except ImportError:
    _DOCX_AVAIL = False


def _extract_doc_text(filename, file_bytes):
    """从上传文件字节中提取纯文本，返回 (text, err)。支持 .pdf / .docx / 纯文本。"""
    name = (filename or "").lower()
    try:
        if name.endswith(".pdf"):
            if not _PDF_AVAIL:
                return "", "服务端未安装 pdfplumber，无法解析 PDF"
            out = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        out.append(t)
            text = "\n\n".join(out).strip()
            if not text:
                return "", "PDF 未提取到文字（可能是扫描件/图片型 PDF，暂不支持 OCR）"
            return text, None
        if name.endswith(".docx"):
            if not _DOCX_AVAIL:
                return "", "服务端未安装 python-docx，无法解析 Word"
            dd = _docx.Document(io.BytesIO(file_bytes))
            paras = [p.text for p in dd.paragraphs if p.text and p.text.strip()]
            for tbl in dd.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                    if cells:
                        paras.append(" | ".join(cells))
            text = "\n\n".join(paras).strip()
            if not text:
                return "", "Word 文档未提取到文字"
            return text, None
        if name.endswith(".doc"):
            return "", "旧版 .doc 二进制格式不支持，请另存为 .docx 或 PDF 后再上传"
        for enc in ("utf-8", "gbk", "latin-1"):
            try:
                return file_bytes.decode(enc).strip(), None
            except Exception:
                continue
        return "", "无法解码文本文件"
    except Exception as e:
        return "", "解析失败：" + str(e)


# 入库噪声过滤：附注/免责声明/页码/落款等模板话不入 RAG
import re as _re_noise
_NOISE_PATTERNS = [
    r"^附注[:：]",
    r"^备注[:：]",
    r"以(正式)?(发文|公告|文件|官方)为准",
    r"^本文档",
    r"^本文件",
    r"^本报告仅",
    r"^仅供(内部)?(参考|学习)",
    r"^免责声明",
    r"^版权所有",
    r"^未经(授权|许可)",
    r"^转载请",
    r"^发文机关[:：]",
    r"^执行时间[:：]",
    r"^印发(日期|时间)[:：]",
]
_NOISE_RE = [_re_noise.compile(p, _re_noise.I) for p in _NOISE_PATTERNS]


def _is_noise_chunk(text):
    """判断一个段落是否为无信息价值的模板/元信息话术。"""
    s = (text or "").strip()
    if not s:
        return True
    for r in _NOISE_RE:
        if r.search(s):
            return True
    # 极短且不含任何数字与汉字的行（孤立标点、页码残片）视为噪声
    if len(s) < 12 and not _re_noise.search(r"[\u4e00-\u9fa5\d]", s):
        return True
    return False


LOG_PATH    = os.path.expanduser("~/.violoop/services/kb-server/rag-audit.log")

def _chunk_text(c):
    """从 chunk 提取纯文本，兼容字符串 / 结构化对象。"""
    if isinstance(c, str):
        return c
    if isinstance(c, dict):
        return str(c.get("text", "") or "")
    return ""

import re as _re
_STOP = set("的了和与及在是有为对以及等就都也很更最这那你我他她它们个之其而或且被把从向到于对由于关于以为")
def _tokenize(s):
    """中英文混合粗分词：英文按词，中文按 2-gram + 单字，去停用词/短词。"""
    s = (s or "").lower()
    toks = set()
    # 英文/数字词
    for w in _re.findall(r"[a-z0-9]+", s):
        if len(w) >= 2:
            toks.add(w)
    # 中文串
    for seg in _re.findall(r"[\u4e00-\u9fff]+", s):
        for ch in seg:
            if ch not in _STOP:
                toks.add(ch)
        for i in range(len(seg) - 1):
            bg = seg[i:i+2]
            if bg[0] not in _STOP or bg[1] not in _STOP:
                toks.add(bg)
    return toks

def _retrieve_chunks(question, chunks, top_k=8, min_score=1):
    """按问题与 chunk 文本的词重合度打分，返回 TopK 真实命中（不足则少返，无命中返空）。
    - 上传/标注材料给予权重加成（origin=admin/user 更权威，应优先吃进）。
    返回: (selected_chunks, scored_debug)"""
    q_toks = _tokenize(question)
    if not q_toks or not chunks:
        # 无法打分（空问题）时退回原顺序前 top_k，保持可用
        return list(chunks)[:top_k], []
    scored = []
    for c in chunks:
        txt = _chunk_text(c)
        c_toks = _tokenize(txt)
        if not c_toks:
            continue
        overlap = q_toks & c_toks
        score = len(overlap)
        # 上传/标注材料加成：命中即 +2，鼓励优先采纳用户提供的证据
        if score > 0 and _is_upload_chunk(c):
            score += 2
        if score >= min_score:
            scored.append((score, c))
    # 按分数降序，稳定保留原相对顺序
    scored.sort(key=lambda x: x[0], reverse=True)
    selected = [c for _s, c in scored[:top_k]]
    return selected, scored[:top_k]

def _is_upload_chunk(c):
    """判断一个知识片段是否来自用户上传/标注（非 AI 采集）。"""
    if not isinstance(c, dict):
        return False
    # 结构化 chunk：origin 为 admin/user 即非 AI 采集
    if c.get("origin") in ("admin", "user"):
        return True
    cid  = str(c.get("id", ""))
    cite = str(c.get("cite", ""))
    tags = c.get("tags", []) or []
    if cid.startswith("file:"):
        return True
    if any(t in ("上传", "材料") for t in tags):
        return True
    # cite 带常见文件扩展名 → 视为上传文件
    if any(cite.lower().endswith(ext) for ext in
           (".txt", ".md", ".csv", ".json", ".pdf", ".doc", ".docx", ".xls", ".xlsx")):
        return True
    return False

def audit_log(entry):
    """把一条 RAG 调用审计记录写入：①本地 JSONL 文件 ②云端 DB(RAG_AUDIT数组，跨会话可查)。"""
    entry["ts"] = datetime.datetime.now().isoformat(timespec="seconds")
    entry["tsMs"] = int(datetime.datetime.now().timestamp() * 1000)
    # ① 本地文件（开发环境）
    try:
        with open(LOG_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps(entry, ensure_ascii=False) + "\n")
    except Exception as e:
        print(f"[audit] file write failed: {e}")
    # ② 云端 DB 持久化：存进同一 store 的 RAG_AUDIT 数组（最近 200 条）
    try:
        if _PG_AVAIL and DATABASE_URL:
            store = json.loads(_db_get() or '{}')
            arr = store.get("RAG_AUDIT") or []
            arr.append(entry)
            store["RAG_AUDIT"] = arr[-200:]
            _db_set(json.dumps(store, ensure_ascii=False))
    except Exception as e:
        print(f"[audit] db write failed: {e}")
    up = entry.get("upload_chunks", 0)
    tot = entry.get("total_chunks", 0)
    print(f"[audit] {entry['ts']} city={entry.get('city','')} mode={entry.get('mode','')} "
          f"chunks={tot} (上传={up}) q=\"{entry.get('question','')[:40]}\"")
    if entry.get("upload_samples"):
        for s in entry["upload_samples"]:
            print(f"[audit]   📎 来自上传《{s['cite']}》: {s['text'][:60]}")

# HTML 与 server.py 同目录（云端部署打包在一起）
SYNC_PATH    = os.environ.get("SYNC_PATH", "/tmp/sync_data.json")
DATABASE_URL = os.environ.get("DATABASE_URL", "")

def _db_conn():
    return psycopg2.connect(DATABASE_URL)

def _init_db():
    if not (_PG_AVAIL and DATABASE_URL):
        print("[db] 未配置 DATABASE_URL，使用文件存储降级模式")
        return
    try:
        conn = _db_conn()
        cur = conn.cursor()
        cur.execute("""
            CREATE TABLE IF NOT EXISTS sync_data (
                id INTEGER PRIMARY KEY DEFAULT 1,
                data TEXT NOT NULL DEFAULT '{}',
                updated_at TIMESTAMP DEFAULT NOW(),
                CHECK (id = 1)
            )
        """)
        conn.commit(); cur.close(); conn.close()
        print("[db] PostgreSQL sync_data 表就绪")
    except Exception as e:
        print(f"[db] 初始化失败: {e}")

def _db_get():
    try:
        conn = _db_conn(); cur = conn.cursor()
        cur.execute("SELECT data FROM sync_data WHERE id=1")
        row = cur.fetchone(); cur.close(); conn.close()
        return row[0] if row else '{}'
    except Exception as e:
        print(f"[db] 读取失败: {e}"); return None

def _db_set(data_str):
    try:
        conn = _db_conn(); cur = conn.cursor()
        cur.execute(
            "INSERT INTO sync_data(id,data,updated_at) VALUES(1,%s,NOW()) "
            "ON CONFLICT(id) DO UPDATE SET data=EXCLUDED.data,updated_at=NOW()",
            (data_str,))
        conn.commit(); cur.close(); conn.close(); return True
    except Exception as e:
        print(f"[db] 写入失败: {e}"); return False

def _keep_nonempty(existing_val, incoming_val):
    """入参为空（None/空串/空 list/dict）时不覆盖已有非空值——防止旧快照把刚同步的数据冲掉。"""
    if incoming_val is None:
        return existing_val
    if isinstance(incoming_val, (list, dict, str)) and len(incoming_val) == 0:
        return existing_val
    return incoming_val

def _kb_material_count(kb):
    """统计一个 project.kb 的累计材料条数（4 主题 known[] 之和）。"""
    if not isinstance(kb, list):
        return 0
    n = 0
    for t in kb:
        if isinstance(t, dict):
            n += len(t.get("known") or [])
    return n

def _merge_map(old, new):
    """按 key 合并 dict-of-dict（PROJECTS / REPORTSTATE / CITY_ACCOUNTS）：
    - 逐字段保留非空值；新增 key 直接并入。
    - kb 字段特判：非空列表也可能是「4 空主题」的占位骨架，按累计材料条数比较，
      材料更多的一方胜出——防止旧快照(0 材料)冲掉刚推送的 RAG。
    - REPORTSTATE 的 text 同理：更长的正文胜出。"""
    merged = {k: (dict(v) if isinstance(v, dict) else v) for k, v in (old or {}).items()}
    for k, v in (new or {}).items():
        if not isinstance(v, dict):
            merged[k] = v; continue
        base = merged.get(k)
        if not isinstance(base, dict):
            merged[k] = dict(v); continue
        for fk, fv in v.items():
            if fk == "kb":
                # 材料多的一方胜出（相等时接受新值，允许内容更新）
                if _kb_material_count(fv) >= _kb_material_count(base.get("kb")):
                    base[fk] = fv
                continue
            if fk == "text":
                if len(fv or "") >= len(base.get("text") or ""):
                    base[fk] = fv
                continue
            base[fk] = _keep_nonempty(base.get(fk), fv)
    return merged

_BASE       = os.path.dirname(os.path.abspath(__file__))
HTML_GOV    = os.path.join(_BASE, "index.html")

# ===== AI 分类+摘要 =====
_KB_SUMMARY_CACHE = {}  # key: md5(city+all_raw) -> {topic: [items]}

SYSTEM_KB_CLASSIFY = """你是招商数据分析师。将以下城市的原始调研材料分类整理到4个板块中，每个板块提炼4-6条核心结论。

四个板块定义：
1. 主导产业与产业链：该城市有哪些主导产业、产业规模、产值、增速、全国地位、核心缺口
2. 园区与承载条件：经济总量(GDP)、园区面积/产能/入驻率、土地/人力/物流成本、基建条件
3. 链主与存量企业：具体龙头企业名称、产能、营收、在建项目、配套率、可对接标的
4. 政策、规划与领导关注：政策文件、补贴力度、竞争城市对比、招商方向优劣势

输出格式（严格遵守，不要额外文字）：
[主导产业与产业链]
结论1
结论2
...
[园区与承载条件]
结论1
结论2
...
[链主与存量企业]
结论1
结论2
...
[政策、规划与领导关注]
结论1
结论2
...

要求：
- 每条结论带关键数字（金额/产量/增速/占比），50-80字
- 覆盖材料中的所有产业方向（不要只挑一个产业）
- 严格按内容归类，不要混淆
- 不要加序号、不要加来源、不要加标题前缀"""

def _classify_and_summarize(city, kb_sections, all_cleaned_items):
    """把一个城市的所有cleaned材料汇总，调用DeepSeek一次性分类+概括到四个板块"""
    if not DS_KEY or not all_cleaned_items:
        return None
    raw_text = '\n'.join(all_cleaned_items)
    cache_key = hashlib.md5((city + raw_text).encode()).hexdigest()
    if cache_key in _KB_SUMMARY_CACHE:
        return _KB_SUMMARY_CACHE[cache_key]
    # 如果所有板块内容都已很精简，跳过
    total_items = len(all_cleaned_items)
    if total_items <= 16 and all(len(x) <= 80 for x in all_cleaned_items):
        return None
    # 调用 DeepSeek 一次性分类+概括
    prompt = f"城市：{city}\n\n原始材料共{total_items}条：\n" + raw_text
    payload = json.dumps({
        "model": MODEL,
        "messages": [
            {"role": "system", "content": SYSTEM_KB_CLASSIFY},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": 1500,
        "temperature": 0.3,
        "stream": False,
        "thinking": {"type": "disabled"}
    }).encode()
    req = urllib.request.Request(
        DS_URL, data=payload, method="POST",
        headers={"Content-Type": "application/json",
                 "Authorization": f"Bearer {DS_KEY}"}
    )
    try:
        with DS_OPENER.open(req, timeout=60) as resp:
            result = json.loads(resp.read())
        content = result.get("choices", [{}])[0].get("message", {}).get("content", "")
        if not content:
            return None
        # 解析分类结果
        classified = {}
        current_topic = None
        topic_map = {
            '主导产业与产业链': '主导产业与产业链',
            '园区与承载条件': '园区与承载条件',
            '链主与存量企业': '链主与存量企业',
            '政策、规划与领导关注': '政策、规划与领导关注',
        }
        for line in content.strip().split('\n'):
            line = line.strip()
            if not line:
                continue
            # 检查是否是板块标题
            for key in topic_map:
                if key in line and '[' in line:
                    current_topic = key
                    break
            else:
                if current_topic and len(line) >= 10:
                    clean_line = line.lstrip('•·-0123456789. ')
                    if len(clean_line) >= 10:
                        classified.setdefault(current_topic, []).append(clean_line)
        if classified:
            _KB_SUMMARY_CACHE[cache_key] = classified
            for t, items in classified.items():
                print(f"[classify] {city}/{t}: {len(items)} items")
            return classified
    except Exception as e:
        print(f"[classify] failed for {city}: {e}")
    return None



def _clean_kb_text(text):
    """清洗 known 条目：删除URL/来源/元信息/系统prompt/章节标题等冗余"""
    if not isinstance(text, str):
        return text
    # 删除所有 URL
    text = _re.sub(r'https?://[^\s\)）,，。]*', '', text)
    # 删除来源标注
    text = _re.sub(r'[\(（]?来源[:：]?\s*[^\)）\n]*[\)）]?', '', text)
    text = _re.sub(r'数据[来源]+[:：][^。；\n]*[。；]?', '', text)
    # 删除日期标注
    text = _re.sub(r'[\(（]\s*\d{4}[-/]\d{2}[-/]\d{2}\s*[\)）]', '', text)
    text = _re.sub(r'[,，]\s*\d{4}[-/]\d{2}[-/]\d{2}', '', text)
    # 删除系统prompt残留
    text = _re.sub(r'编制单位[:：][^。\n]*', '', text)
    text = _re.sub(r'数据基准[:：][^。\n]*', '', text)
    text = _re.sub(r'引用铁律[:：][^。\n]*', '', text)
    text = _re.sub(r'地价只认[^。\n]*[。]?', '', text)
    text = _re.sub(r'政策必标[^。\n]*[。]?', '', text)
    text = _re.sub(r'>\s*[^\n]*', '', text)
    # 删除所有【...】及孤立的】
    text = _re.sub(r'【[^】]*】\s*', '', text)
    text = text.replace('】', '')
    # 删除生成时间/分析对象等元信息
    text = _re.sub(r'生成时间[:：][^。；\n]*[。；]?', '', text)
    text = _re.sub(r'分析对象[:：][^。；\n]*[。；]?', '', text)
    text = _re.sub(r'gov\.cn\s*权威源为主[）\)]?', '', text)
    # 删除文号标注 如 （鄂政发〔2021〕29 号）
    text = _re.sub(r'[\(（][^\)）]*〔\d+〕\d+\s*号[\)）]', '', text)
    # 清理多余空格/标点
    text = _re.sub(r'\s{2,}', ' ', text)
    text = _re.sub(r'[。；，]{2,}', '。', text)
    # 清理 markdown 粗体标记（所有**都删掉）
    text = text.replace('**', '')
    # 清理 URL 路径碎片
    text = _re.sub(r'\S*\.s?html\S*', '', text)
    text = _re.sub(r'\S*/\d{8}/[a-f0-9]+/\S*', '', text)
    text = text.strip(' 。；，')
    return text

def _is_junk_item(text):
    """判断一条 known 是否为无用条目（纯标题/元信息/报告头）应被丢弃"""
    if not text or len(text) < 6:
        return True
    # 纯报告标题类（不含有用数据）
    junk_patterns = [
        r'^[✅⚠️\s]*[\[【]?.*招商方向研判报告[】\]]?[。\s]*$',
        r'^[✅⚠️\s]*[\[【]?.*经济数据报告[】\]]?.*$',
        r'^[✅⚠️\s]*[\[【]?.*竞争格局分析报告[】\]]?.*$',
        r'^[✅⚠️\s]*[\[【]?.*招商作战清单[】\]]?.*$',
        r'^[✅⚠️\s]*[\[【]?.*市方向[一二三四五六七八九十\d]+.*$',
        r'^[✅\s]*\d+年各产业链方向.*请领导确认$',
        r'^[✅\s]*(首选|拟优先|专项资金).*请(领导|主管).*确认$',
    ]
    for pat in junk_patterns:
        if _re.search(pat, text):
            return True
    # 纯表格行/URL碎片/分隔线
    if _re.match(r'^\d{3,4}/t\d+', text):
        return True
    if _re.match(r'^注[:：]', text):
        return True
    if _re.match(r'^(年份|排名)\s+(GDP|企业)', text):
        return True
    if '------' in text or '---' * 3 in text:
        return True
    if _re.match(r'^\d+\s+\S+\s+★', text):
        return True
    # URL碎片
    if '.shtml' in text or '.html)' in text:
        return True
    # fortune/xx 类残留
    if _re.search(r'ortune/\d+/', text):
        return True
    # 纯数字串
    if _re.match(r'^[\d\s\.]+$', text.replace(',', '').replace('，', '')):
        return True
    # 趋势判断框架
    if text.startswith('趋势判断'):
        return True
    # 含★排名的企业表格行
    if '★★' in text:
        return True
    # 数据残片（如"数据2025年结构比"）
    if _re.match(r'^数据\d{4}', text):
        return True
    return False

def _clean_sync_data(data):
    """直接返回原始数据（紧急恢复，不做任何修改）"""
    return data

HTML_OPS    = os.path.join(_BASE, "ops.html")
DS_URL      = "https://api.deepseek.com/v1/chat/completions"
# 直连 DeepSeek，绕过系统代理(Clash 7897)——否则请求会挂死
DS_OPENER   = urllib.request.build_opener(urllib.request.ProxyHandler({}))
def _load_ds_key():
    # 优先精确匹配环境变量；找不到则容错扫描（变量名含空格/DEEPSEEK/DS_KEY 都能兜住）
    v = os.environ.get("DEEPSEEK_API_KEY", "")
    if not v.strip():
        for k, val in os.environ.items():
            ku = k.strip().upper().replace(" ", "").replace("-", "_")
            if ku in ("DEEPSEEK_API_KEY", "DEEPSEEKAPIKEY", "DS_KEY", "DEEPSEEK_KEY") or "DEEPSEEK" in ku:
                if val and val.strip():
                    v = val; break
    # key 仅从环境变量读取（Railway Variables 配置 DEEPSEEK_API_KEY）；不再从明文文件兜底，避免密钥入库泄露
    v = v.strip().strip('"').strip("'").strip()
    if v and not v.startswith("sk-"):
        import re as _re
        m = _re.search(r"sk-[A-Za-z0-9_\-]{10,}", v)
        if m: v = m.group(0)
    return v
DS_KEY      = _load_ds_key()
MODEL       = "deepseek-v4-pro"
# chat 模式走 Responses API + 原生联网搜索（仅 v4-flash 支持 web_search，pro 暂不支持）
DS_RESP_URL = "https://api.deepseek.com/responses"
MODEL_CHAT  = "deepseek-v4-flash"
MAX_TOKENS_DRAFT = 800
MAX_TOKENS_FULL  = 5000
MAX_TOKENS_CHAT  = 400
# research 模式需返回结构化 JSON（对标企业3~5家+数量+来源），且 v4-pro 默认 thinking
# 会消耗 token，故单独放大并预留思考空间
MAX_TOKENS_RESEARCH = 6000
# 整链研判：一次返回整条链所有环节（6~8个），需更大 token 预算
MAX_TOKENS_CHAIN = 12000
# 云端：单端口读 $PORT（Railway 注入）；本地默认 5050。政府端/管理端同端口，ops 走 /ops 路由
PORT_GOV    = int(os.environ.get("PORT", "5050"))
PORT_OPS    = 5051

SYSTEM_DRAFT = """你是慧小招AI招商助手。快速输出300字以内的初步分析草稿：
①关键缺口（附数据）②招引方向（标类型）③待确认事项（每条⚠️开头）④数据可靠性。
用⚠️标注不确定判断。语言精炼，禁止五章报告。"""

SYSTEM_CHAT = """你是慧小招城市智库的AI问答助手，在招商问答对话框中与政府干部自然对话。

要求：
- 像聊天一样简短直接地回答，控制在150字以内。
- 只回答用户当前问的这一个问题，不要输出报告结构、不要分章节、不要生成表格。
- 优先引用提供的城市智库数据与用户上传材料中的具体数字。
- 当问题涉及最新动态、时效信息（如近期政策、行业新闻、企业动向）或智库数据无法覆盖时，使用联网搜索获取最新信息，并注明「（联网）」。
- 无数据支撑且未联网核实的判断用⚠️标注。
- 如果用户只是补充/更新了一条数据，简短确认已收到并说明它对研判的意义即可，不要展开长篇分析。
- 使用中文，口语化、精炼。"""

SYSTEM_SUGGEST = """你是慧小招城市智库的AI助手。根据提供的某个城市的智库数据，站在当地招商干部视角，生成他们此刻最该问、最有价值的建议问题。

要求：
- 只输出 4 个问题，每行一个，不加序号、不加符号、不加任何解释。
- 每个问题必须紧扣该城市的真实产业/园区/链主/政策特征，带上城市名或该城市的具体产业名，禁止泛泛而谈。
- 问题要具体、可回答、对招商决策有用（围绕补链缺口、承接方向、园区分工、竞争差异化、招引对象等）。
- 若智库数据不足，也要基于城市名和常识提出该城市可能关心的招商问题，禁止出现其他城市的名字。
- 使用中文，每个问题不超过25字。"""

SYSTEM_ONBOARD = """你是慧小招招商顾问。为某城市的招商偏好问卷生成\u201c结合本地实情\u201d的推荐选项，帮政府招商干部快速作答。

给你城市名与该市城市智库知识片段。请结合两者（智库片段优先，不足时用你对该城市的了解补充）为下面 5 个问题各生成推荐选项：
- park: 该市重点招商的产业园区（尽量给真实园区名称，如\u201cXX高新区\u201d\u201cXX经开区\u201d）
- capacity: 该市产业园区的承载能力（厂房/用地供给现状）
- scale: 适合该市的理想招商企业规模
- industry: 该市优先招引的产业方向（结合本地主导/培育产业，给真实方向名）
- invest: 该市期望的企业投资强度

严格要求：
- 只输出一个 JSON 对象，不要任何解释、不要 markdown 代码块围栏。
- 结构：{"park":["选项1","选项2",...],"capacity":[...],"scale":[...],"industry":[...],"invest":[...]}
- 每题 3-5 个选项，每个选项不超过 18 字，具体、贴合该城市，禁止出现其他城市名。
- park 与 industry 必须尽量给该市真实的园区名/产业名；其余题结合该市体量给合理档位。
- 若完全没有该城市信息，也要基于城市名和常识给出合理推荐，绝不留空。"""

SYSTEM_CITY_PROFILE = """你是慧小招城市画像分析师。基于该城市智库(RAG)的全量材料，以及可能存在的招商偏好问卷answers，刻画这个城市的①客观自身条件 ②招商偏好。

## 铁律
- 只能使用给定知识片段中的事实。每一条结论必须能在片段中找到依据。
- 客观条件必须带具体数字/企业名/园区名；没有数据的字段留空字符串，绝不编造、绝不用"约""预计"糊过去。
- 招商偏好分两类来源，必须区分标注：
  · stated = 问卷明确作答 或 干部访谈明确表达的（confidence 高）
  · inferred = 从产业基础/报告方向/政策材料反推出来的（confidence 中/低，必须在 basis 写清反推依据）
- 若问卷与反推结论冲突，以问卷为准，并在 conflicts 中记录该冲突。
- 访谈中"不知道""很一般""没有"这类未提供实质信息的答复，不得当作实质偏好，应记入 engagement。
- **engagement 措辞铁律**（画像要给客户看，不得出现指责对方的词）：
  · 只陈述「哪些信息尚未获取」，绝不评价人的态度、意愿或配合程度。
  · 禁用词：配合度、敷衍、回避、不配合、消极、推诿。
  · 不要原话引用干部的答复（如"你自己去想办法"），改写成中立的信息缺口描述。
  · 正例："产业基础现状、专项政策额度、已接触目标企业三项访谈尚未取得具体信息，建议下轮对接补充"。
  · 反例（禁止）："干部答不知道、很一般，配合度低"。
  · level 只反映信息完备度：high=较完备 / medium=部分完备 / low=待补充。

## 输出（严格只输出一个 JSON 对象，不要 markdown 围栏、不要任何解释文字）
{
  "summary": "150字内一段话总结这个城市是什么样的招商标的：底子、长板、短板、最该招什么",
  "objective": {
    "economy":   [{"label":"GDP总量","value":"1502.81亿元(2025)","cite":"片段来源标签"}],
    "structure": [{"label":"三次产业结构","value":"14.1:38.5:47.4","cite":""}],
    "industry":  [{"label":"主导产业","value":"专用汽车","detail":"全国占比/地位说明","cite":""}],
    "park":      [{"label":"园区名称","value":"承载能力/厂房/用地现状","cite":""}],
    "anchor":    [{"label":"链主企业","value":"产能/营收/在建","cite":""}],
    "cost":      [{"label":"配套率","value":"41%","detail":"对比城市","cite":""}]
  },
  "preference": {
    "industry": {"items":["优先产业方向"],"source":"stated|inferred","confidence":"high|medium|low","basis":"依据说明"},
    "park":     {"items":[],"source":"","confidence":"","basis":""},
    "scale":    {"items":[],"source":"","confidence":"","basis":""},
    "invest":   {"items":[],"source":"","confidence":"","basis":""},
    "capacity": {"items":[],"source":"","confidence":"","basis":""}
  },
  "strength":  ["不可复制的优势，带数据"],
  "weakness":  ["真实短板/缺口，带数据"],
  "conflicts": ["问卷与材料反推之间的矛盾点，没有则空数组"],
  "engagement":{"level":"high|medium|low","signals":["尚未取得具体信息的访谈环节（中立表述，不评价态度）"]},
  "gaps":      ["画像缺失项：还需要补什么材料才能把画像做实"],
  "charts": {
    "structure":  {"primary":14.1,"secondary":38.5,"tertiary":47.4,"year":"2024","cite":""},
    "gdp_trend":  [{"year":"2020","value":1096.72},{"year":"2024","value":1442.35}],
    "growth":     [{"label":"GDP增速","value":6.1,"unit":"%"},{"label":"规上工业增加值","value":9.9,"unit":"%"}],
    "chain":      [{"node":"整车改装","status":"strong"},{"node":"动力电池电芯","status":"missing"}],
    "compare":    {"metric":"配套率","self":{"label":"随州","value":41},
                   "peers":[{"label":"十堰","value":75},{"label":"梁山","value":65}],"unit":"%"}
  }
}

## charts 专项规则（这是纯图表区的数据源，必须尽力填满）
- 只填能从材料里拿到**纯数字**的项；拿不到的整个键省略（不要填 0、不要填 null、不要猜）。
- structure：三次产业占比，三个数加起来应≈100。取最新年份。
- gdp_trend：GDP 总量时间序列，2-6 个点，按年份升序，value 单位统一为「亿元」。只有一个年份的数据就省略此键。
- growth：增速类指标 2-6 条，value 为纯数字（% 值写 6.1 不写 "6.1%"）。
- chain：产业链环节强弱，4-10 个节点。status 只能取 strong(本地强)/weak(薄弱)/missing(缺失)。
  这是招商最关心的图——务必从"缺口/短板/配套"类材料里提取。
- compare：本市 vs 竞争城市的同一可量化指标（如配套率、企业数、产值）。peers 1-4 个。
  self 和 peers 必须是同一口径同一单位，否则省略此键。

## 要求
- objective 每个分组 0-6 条，有几条写几条，宁缺勿造。
- cite 填知识片段前的来源标签（形如 [来源·主题]）中的可辨识信息，无法确定填空字符串。
- preference 五个维度都要出现；完全无依据的维度 items 留空数组、confidence 填 "low"、basis 说明为何无法判断。
- gaps 要具体可执行（如"缺2025年园区可供地块四至与亩数"），不要写"信息不足"。"""

SYSTEM_FULL = """你是慧小招产业链研判师，结合城市智库数据完成完整的招商研判报告。

## 报告结构（严格按此输出）

### 一、产业基础判断
基于知识片段，描述城市产业现状（产值/规模/配套率等具体数字）。

### 二、产业链缺口分析（核心）
对每个研判方向，逐环节标注状态：
- ✅ 已有：[企业名] 承担 [环节]，规模/能力说明
- ⚠️ 薄弱：[环节]，现状说明，不足在哪
- ❌ 缺失：[环节]，占整体成本/价值XX%，全部外购自[地区/企业]

对缺失/薄弱环节标注本地化属性：
A=必须本地化（依赖原料/链主/物流）
B=可跨区域（技术密集型）
C=优先本地化（有协同但不硬依赖）

### 三、补链优先级清单（TOP 5）
| 排名 | 缺口节点 | 本地化属性 | 经济拉动★ | 招引可行性★ | 综合优先级 |

### 四、目标企业画像（每个优先缺口）
对每个TOP缺口描述：
- 目标企业类型（规模/技术路线/上市/非上市）
- 核心诉求（为什么要来这里）
- 开口话术：「[城市]的[不可复制资产]正是您最需要的[企业具体需求]，落地即能锁定[链主名]的[XX亿]采购订单。」

### 五、待确认事项
列出需要确认的关键事项（专项资金/园区地块/政策口径）。
⚠️ 标注每条待确认项。

要求：引用具体数字支撑每个判断；企业名和数据来自知识片段；无数据支撑的判断用⚠️标注。

## 企业推荐严格标准（必须同时满足）
1. 补链匹配：企业产品直接填补报告中标注的❌缺失环节
2. 扩张信号（必须有其中一项可查证信号）：
   - 新建/扩建工厂的公开公告或政府签约
   - 在目标区域的招聘岗位或调研记录
   - 融资公告明确注明产能扩张用途
   - 与当地政府/园区签署合作备忘录
   ⚠️ 仅因企业优秀/行业领先但无扩张信号，禁止列入候选
3. 信号来源：每条必须标注（公告编号/媒体名称/日期）；
   查不到来源须标注「（未获公开来源，待核实）」，禁止虚构

## 数字可靠性铁律（违反即为重大错误）
1. 报告里出现的每一个「数字」（产值/规模/企业数量/占比/金额）都必须能追溯到来源，
   来源限：政府公报 / 国家统计局 / 行业协会白皮书 / 上市公司财报·公告 / 权威媒体 / 工商注册数据。
2. 引用数字时在括号内标注来源与年份，如「(随州市统计局·2024)」「(中国食用菌协会·2025)」。
3. 拿不到可靠来源的数字，一律写「待核实」并说明原因，禁止编造、禁止用"大约/约/估计"糊弄一个看似合理的数。
4. 尤其「全国有多少家企业」「本地多少家企业」这类精确计数，除非有权威统计口径，否则必须标「待核实」。
5. 知识片段中已有的数字可直接引用，但也要能对应到片段标注的来源（cite）。"""

SYSTEM_RESEARCH = """你是慧小招产业调研员，负责针对某个产业链「环节」做完整调研：①该环节在本地的强弱判定（优势/培育/缺口）②本地代表企业 ③全国对标企业 ④企业数量与市场规模。

## 输出格式（严格输出 JSON，不要 markdown 代码块、不要任何 JSON 之外的文字）
{
  "segment": "环节名",
  "level": "strong|mid|weak",
  "level_basis": "判定依据（一句话，必须标注来源，如：随州为全球白花菇最大产地，约占全国50%——中国食用菌协会2024）",
  "local_leaders": [
    {"name": "本地代表企业名", "role": "它在本地该环节中的角色", "source": "信息来源"}
  ],
  "benchmarks": [
    {"name": "外地对标企业名", "region": "总部城市", "kind": "主营业务(简短)",
     "match": "为何匹配该环节", "signal": "扩张/迁移信号(无则空)", "source": "信息来源"}
  ],
  "local_count": {"value": "本地企业数或null", "source": "来源或null", "note": "无法给出时的原因"},
  "national_count": {"value": "全国企业数或null", "source": "来源或null", "note": "无法给出时的原因"},
  "market_size": {"value": "市场规模或null", "source": "来源或null", "note": "无法给出时的原因"},
  "confidence": "high|medium|low",
  "caveats": ["需人工核实的事项"]
}

## level 判定标准（必须基于可查证事实，不能拍脑袋）
- strong（优势）：本地已有全国知名龙头/产业集群，或产量/规模居全国前列（须有权威来源佐证）
- mid（培育）：本地已有一定基础，但规模/实力不足，需培育或增强
- weak（缺口）：本地该环节缺失、或严重依赖外购（须有"缺失/外购"的可查证事实）

## 铁律（违反即为重大错误）
1. level 判定必须有 level_basis 依据 + 来源；拿不准就判 mid 并在 caveats 标"待核实"。
2. 每一个数字（企业数量/市场规模）都必须有可靠来源；拿不到就 value=null + note"无权威公开口径，待核实"，禁止编造、禁止给一个看似合理的估算值。
3. local_leaders 只列本地（该城市/下辖区县）真实存在的企业；查不到本地企业就返回空数组 []，禁止拿外地企业冒充。
4. benchmarks 只列外地（非本地）真实对标企业，每家有 source；查不到就不列。
5. 企业必须真实存在，禁止虚构企业名；宁可少列，不可编造。"""

SYSTEM_RESEARCH_CHAIN = """你是慧小招产业调研员，负责对某城市一条完整产业链的「所有环节」一次性做研判。每个环节都要给出：强弱判定（优势/培育/缺口）+ 判定依据 + 本地代表企业 + 外地对标企业。

## 输出格式（严格输出 JSON，不要 markdown 代码块、不要任何 JSON 之外的文字）
{
  "chain": "产业链名",
  "segments": [
    {
      "name": "环节名",
      "level": "strong|mid|weak",
      "level_basis": "判定依据（一句话，必须标注来源与年份）",
      "local_leaders": [{"name": "本地代表企业名", "role": "角色", "source": "来源"}],
      "benchmarks": [{"name": "外地对标企业名", "region": "总部城市", "kind": "主营业务",
                      "match": "为何匹配", "signal": "扩张/迁移信号(无则空)", "source": "来源"}],
      "local_count": {"value": "本地企业数或null", "source": "来源或null"},
      "national_count": {"value": "全国企业数或null", "source": "来源或null"}
    }
  ],
  "confidence": "high|medium|low",
  "caveats": ["需人工核实的事项"]
}

## level 判定标准（必须基于可查证事实，不能拍脑袋）
- strong（优势）：本地已有全国知名龙头/产业集群，或产量/规模居全国前列（须有权威来源佐证）
- mid（培育）：本地已有一定基础，但规模/实力不足，需培育或增强
- weak（缺口）：本地该环节缺失、或严重依赖外购（须有"缺失/外购"的可查证事实）

## 铁律（违反即为重大错误）
1. 每个环节的 level 判定必须有 level_basis 依据 + 来源；拿不准就判 mid 并在 caveats 标"待核实"。
2. 每一个数字（企业数量 local_count/national_count）都必须有可查证的权威来源；拿不到就 value=null 且 source=null，禁止编造、禁止给一个看似合理的估算值。
3. local_leaders：列出你「确定真实存在」的本地（该城市/下辖区县）企业，优先列知名龙头/上市公司/出口企业/行业标杆，source 写「公开工商信息」即可。企业是否真实存在以你的可靠知识为准，不要因为"无法实时联网核实"就空着——只有该环节确实没有已知本地企业时才返回 []。禁止拿外地企业冒充本地企业。
4. benchmarks：只列外地（非本地）真实对标企业，优先列上市公司/细分龙头，每家有 source（公开资料即可）；确实没有就不列。
5. 企业名必须真实存在、用规范全称（如「湖北裕国菇业股份有限公司」「品源（随州）现代农业发展有限公司」），禁止虚构企业名。
6. 环节名必须与用户给出的环节清单完全一致，一个不少，不要增删改。"""


SYSTEM_TOPICS = """你是慧小招产业招商策略师。根据某城市的调研报告（知识片段），归纳出该城市当前最值得推进的 3-5 个「产业招商分析方向」。方向要具体、可招商落地，通常是「补链缺口」「精深加工升级」「承接转移」这类可执行主题，而不是宽泛的产业名。

## 输出格式（严格输出 JSON，不要 markdown 代码块、不要 JSON 之外的任何文字）
{
  "topics": [
    {
      "label": "方向名称（8-16字，具体可招商，如「氢能专用车核心零部件补链」）",
      "icon": "一个贴切的emoji",
      "type": "补链|精深加工|承接转移|培育",
      "desc": "一句话依据（≤24字，必须来自知识片段的事实，如「电堆占成本53%全外购」）"
    }
  ]
}

## 铁律
1. 每个方向的 label 和 desc 必须能在知识片段里找到事实支撑；片段没提到的产业，不要凭空造方向。
2. desc 优先引用知识片段里的具体数字/事实（占比、产值、企业数、缺口环节），不要空话套话。
3. 输出 3-5 个方向，按招商价值/紧迫性从高到低排序（缺口大、拉动强的排前面）。
4. 只输出该城市真实相关的方向；宁可少，不可为凑数编造。
5. icon 用单个 emoji；type 从给定四类里选最贴切的。"""


SYSTEM_FUNNEL = """你是慧小招招商情报分析师。针对某个具体的「招商项目方向」，基于其产业研判报告与缺口分析，筛选出真实存在的、有异地设厂/投资/合作可能的适配企业，构成一个企业漏斗。

## 输出格式（严格输出 JSON，不要 markdown 代码块、不要 JSON 之外的任何文字）
{
  "topic": "该招商方向名",
  "total_scanned": 40,
  "companies": [
    {
      "name": "企业规范全称（真实存在的公司，如「亿华通科技股份有限公司」）",
      "region": "总部所在城市",
      "kind": "主营业务（简短，一句话）",
      "fit": "为何适配该招商方向（结合缺口，一句话）",
      "score_match": 36,
      "score_relocate": 22,
      "score_strength": 25,
      "score_reason": "三项打分的简要依据（一句话说明为何这样给分）",
      "signal": "扩张/迁移/投资信号（有则写，无则空字符串）",
      "expansion": "该企业明确的扩张/建厂/异地投资/产能扩产等公开需求信号（务必有公开依据才写，注明信息线索；无则留空字符串）",
      "faction": "该企业核心领导（董事长/总经理/创始人）与本招商城市或所在省是否存在可考的『同派系』关联——校友（同一母校）、同乡（籍贯）、商会/行业协会共同任职、过往任职交集等（务必有可信线索才写，写明关联类型与具体依据，如『董事长张三为武汉大学校友』；无则留空字符串）",
      "listed": "上市情况（如 A股/新三板/未上市/港股）",
      "source": "信息来源（如 公开工商信息 / 上市公司公告 / 行业协会名录）"
    }
  ]
}

## 要求
1. companies 输出目标 40 家（至少 30 家），聚焦该方向缺口环节适配度最高的真实企业；宁可精不可滥，不要为凑数拉低质量。
2. 企业必须是你「确定真实存在」的公司，用规范全称，优先列上市公司/细分龙头/专精特新/已有扩张信号的企业。禁止虚构企业名——宁可少列，绝不编造。
3. 【评分卡·三个分项，分别独立打分，禁止都给高分，要有区分度】
   - score_match（0-40）：业务与该方向缺口环节的匹配度。越是精准补上报告识别出的缺口环节，分越高。
   - score_relocate（0-30）：异地投资/迁移/落地本城市的可能性。有公开扩张/迁移/投资动向或处于产能扩张期的，分越高；总部已在本地或明显无外迁可能的，分低。
   - score_strength（0-30）：企业实力。上市公司/细分龙头/专精特新/规模大的分高。
   - 不要输出 fit_score 总分，总分由系统按三项相加得出。score_reason 用一句话说明三项打分依据。
4. fit 必须具体说明该企业能补上这个方向的哪个缺口环节，不要空话。
5. signal / expansion 只写你确有公开依据的动向；没有就留空字符串，【严禁编造】任何融资/建厂/扩张消息。
6. faction 是稀缺信息：只有当你确有可信线索（公开可考的校友/籍贯/商会/任职关联）时才填，并写清依据；绝大多数企业此项应留空——【严禁为了凑数编造领导背景或人际关联】。找不到就留空，这是正常的，不强求。
7. total_scanned 写你本轮实际扫描评估的企业规模（约等于 companies 长度或更多）。
8. 全部用中文。这是 AI 推理结果，前端会标注"待人工核验"。核心底线：所有企业、分数、信号、扩张需求、派系关联都必须基于真实公开信息或合理推断，绝不编造虚假数据；宁可留空、宁可少列，也不许造假。"""


class Handler(BaseHTTPRequestHandler):
    def log_message(self, fmt, *args):
        print(f"[kb] {self.address_string()} {fmt % args}")

    def cors(self):
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET,POST,OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")

    def do_OPTIONS(self):
        self.send_response(200); self.cors(); self.end_headers()

    def do_GET(self):
        # /ops 和 /ops.html 路径在两个端口都服务管理端
        # 5051 访问 / 也直接服务管理端（和 5050/ops 共享 origin→不行，但5050/ops 共享 origin 可以）
        path = self.path.split('?')[0]
        # 云端数据同步：GET /api/sync 直接返回原始数据
        if path == '/api/sync':
            try:
                if _PG_AVAIL and DATABASE_URL:
                    result = _db_get()
                    raw_str = result or '{}'
                else:
                    try:
                        with open(SYNC_PATH, encoding='utf-8') as f2:
                            raw_str = f2.read()
                    except FileNotFoundError:
                        raw_str = '{}'
                data = raw_str.encode()
            except Exception as e:
                data = json.dumps({"error": str(e)}).encode()
            self.send_response(200)
            self.send_header('Content-Type', 'application/json; charset=utf-8')
            self.send_header('Content-Length', str(len(data)))
            self.cors(); self.end_headers(); self.wfile.write(data)
            return
        # 政府端城市智库AI精选概括（只读，不修改数据库）
        if path == '/api/kb-summary':
            try:
                # 读取原始数据
                if _PG_AVAIL and DATABASE_URL:
                    result = _db_get()
                    raw_str = result or '{}'
                else:
                    try:
                        with open(SYNC_PATH, encoding='utf-8') as f2:
                            raw_str = f2.read()
                    except FileNotFoundError:
                        raw_str = '{}'
                sync_obj = json.loads(raw_str)
                projects = sync_obj.get('PROJECTS') or {}
                cur_proj = sync_obj.get('cur', '')
                proj = projects.get(cur_proj) or {}
                city = proj.get('city', '')
                kb = proj.get('kb') or []
                if not kb or not city:
                    data = json.dumps({"ok": False, "error": "no kb data"}).encode()
                else:
                    # 收集所有板块的cleaned items
                    all_cleaned = []
                    for section in kb:
                        known = section.get('known') or []
                        for item in known:
                            t = _clean_kb_text(item)
                            if not _is_junk_item(t):
                                all_cleaned.append(t)
                    # AI分类+概括（有缓存）
                    classified = _classify_and_summarize(city, kb, all_cleaned)
                    # 构建返回结果
                    result_kb = []
                    for section in kb:
                        topic = section.get('t', '')
                        if classified and topic in classified:
                            items = classified[topic]
                        else:
                            # AI未返回时，用清洗后的前6条
                            cleaned = []
                            for item in (section.get('known') or []):
                                t = _clean_kb_text(item)
                                if not _is_junk_item(t) and len(t) > 10:
                                    cleaned.append(t)
                            items = cleaned[:6]
                        # 修正 sub/tag
                        sub = section.get('sub', '')
                        if 'AI流水线' in sub or '流水线产出' in sub or not sub:
                            _sub_map = {
                                '主导产业与产业链': city + '主导产业集群与核心缺口',
                                '园区与承载条件': city + '主要园区与承载能力',
                                '链主与存量企业': city + '链主企业与配套格局',
                                '政策、规划与领导关注': city + '政策方向与竞争态势',
                            }
                            sub = _sub_map.get(topic, city + '产业数据')
                        tag = section.get('tag', '')
                        if tag in ('已初始化', '') or 'AI流水线' in tag:
                            n = len(items)
                            tag = f'AI研判 · {n}条' if n else '待补充'
                        result_kb.append({
                            "t": topic,
                            "icon": section.get("icon", ""),
                            "sub": sub,
                            "tag": tag,
                            "known": items
                        })
                    data = json.dumps({"ok": True, "city": city, "kb": result_kb}, ensure_ascii=False).encode()
            except Exception as e:
                import traceback; traceback.print_exc()
                data = json.dumps({"ok": False, "error": str(e)}).encode()
            self.send_response(200)
            self.send_header('Content-Type', 'application/json; charset=utf-8')
            self.send_header('Content-Length', str(len(data)))
            self.cors(); self.end_headers(); self.wfile.write(data)
            return
        # 原始报告下载：GET /api/report-file?city=随州&kind=full|short
        # 从云端存储的 base64 docx 里取出并回传，供管理端「下载原始报告」按钮使用。
        if path == '/api/report-file':
            import base64, urllib.parse
            qs = urllib.parse.parse_qs(self.path.split('?', 1)[1] if '?' in self.path else '')
            city = (qs.get('city') or [''])[0]
            kind = (qs.get('kind') or ['full'])[0]
            try:
                if _PG_AVAIL and DATABASE_URL:
                    store = json.loads(_db_get() or '{}')
                else:
                    with open(SYNC_PATH, encoding='utf-8') as f2:
                        store = json.loads(f2.read())
            except Exception:
                store = {}
            # 优先从 REPORT_REQUESTS（最近一条该城市 done 且带 files 的），回退到项目 reportFiles
            files = None
            done = [r for r in (store.get('REPORT_REQUESTS') or [])
                    if isinstance(r, dict) and r.get('city') == city and r.get('files')]
            if done:
                files = done[-1]['files']
            if not files:
                for pv in (store.get('PROJECTS') or {}).values():
                    if isinstance(pv, dict) and pv.get('city') == city and pv.get('reportFiles'):
                        files = pv['reportFiles']; break
            meta = None
            if files:
                meta = next((m for m in files if m.get('kind') == kind), None) or files[0]
            if not meta or not meta.get('b64'):
                msg = json.dumps({'ok': False, 'error': '未找到该城市原始报告文件'}).encode()
                self.send_response(404); self.send_header('Content-Type', 'application/json; charset=utf-8')
                self.send_header('Content-Length', str(len(msg)))
                self.cors(); self.end_headers(); self.wfile.write(msg); return
            try:
                blob = base64.b64decode(meta['b64'])
            except Exception:
                blob = b''
            fname = meta.get('name') or (city + '_报告.docx')
            fname_enc = urllib.parse.quote(fname)
            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition',
                             "attachment; filename*=UTF-8''" + fname_enc)
            self.send_header('Content-Length', str(len(blob)))
            self.cors(); self.end_headers(); self.wfile.write(blob)
            return
        port = self.server.server_address[1]
        if path in ("/", "/index.html"):
            # 本地 5051 端口访问 / → 重定向到同 origin 的 /ops（云端单端口不触发）
            if port == PORT_OPS:
                self.send_response(302)
                self.send_header("Location", "/ops")
                self.cors(); self.end_headers(); return
            html_path = HTML_GOV
        elif path in ("/ops", "/ops.html"):
            html_path = HTML_OPS  # 从 5050/ops 提供，共享 origin
        elif path == "/rag-log":
            # RAG 审计日志：JSON（?format=json）或纯文本（默认）
            fmt = "json" if "format=json" in (self.path.split("?")[1] if "?" in self.path else "") else "text"
            recs = []
            # ① 云端 DB 优先（生产环境跨会话可查）
            try:
                if _PG_AVAIL and DATABASE_URL:
                    store = json.loads(_db_get() or '{}')
                    recs = store.get("RAG_AUDIT") or []
            except Exception as e:
                print(f"[rag-log] db read failed: {e}")
            # ② 回退本地文件（开发环境）
            if not recs:
                try:
                    lines = open(LOG_PATH, encoding="utf-8").read().splitlines()
                    recs = [json.loads(l) for l in lines if l.strip()]
                except FileNotFoundError:
                    recs = []
            if fmt == "json":
                out = json.dumps(recs, ensure_ascii=False, indent=2).encode()
                ctype = "application/json; charset=utf-8"
            else:
                rows = []
                tot_up = sum(r.get("upload_chunks", 0) for r in recs)
                rows.append(f"慧小招 RAG 审计日志 — 共 {len(recs)} 次 AI 分析调用，累计吃进上传片段 {tot_up} 个")
                rows.append("=" * 72)
                for r in recs:
                    rows.append(f"[{r.get('ts','')}] 城市={r.get('city','')} 模式={r.get('mode','')}")
                    rows.append(f"  问题: {r.get('question','')}")
                    rows.append(f"  数据块: 共{r.get('total_chunks',0)} 个 "
                                f"(智库{r.get('kb_chunks',0)} + 上传{r.get('upload_chunks',0)})")
                    rows.append(f"  来源: {'、'.join(r.get('all_cites',[]))}")
                    for s in r.get("upload_samples", []):
                        rows.append(f"    📎 来自上传《{s.get('cite','')}》: {s.get('text','')[:80]}")
                    rows.append("-" * 72)
                out = ("\n".join(rows) + "\n").encode()
                ctype = "text/plain; charset=utf-8"
            self.send_response(200)
            self.send_header("Content-Type", ctype)
            self.send_header("Content-Length", str(len(out)))
            self.cors(); self.end_headers(); self.wfile.write(out)
            return
        else:
            if path == "/health":
                pass  # handled below
            else:
                self.send_error(404); return
        if path not in ("/health",):
            try:
                data = open(html_path, "rb").read()
                self.send_response(200)
                self.send_header("Content-Type", "text/html; charset=utf-8")
                self.send_header("Cache-Control", "no-cache, no-store, must-revalidate")
                self.send_header("Content-Length", str(len(data)))
                self.cors(); self.end_headers(); self.wfile.write(data)
            except FileNotFoundError:
                self.send_error(404, "HTML not found")
        elif self.path == "/health":
            port = self.server.server_address[1]
            role = "ops" if port == PORT_OPS else "gov"
            body = json.dumps({"ok": True, "model": MODEL, "key": bool(DS_KEY),
                               "key_len": len(DS_KEY),
                               "key_prefix": (DS_KEY[:5]+"…"+DS_KEY[-3:]) if DS_KEY else "",
                               "role": role, "port": port}).encode()
            self.send_response(200)
            self.send_header("Content-Type", "application/json")
            self.cors(); self.end_headers(); self.wfile.write(body)
        else:
            self.send_error(404)

    def do_POST(self):
        length = int(self.headers.get('Content-Length', 0))
        raw = self.rfile.read(length)
        # 云端数据同步：POST /api/sync 合并保存（保护字段不被空值覆盖）
        if self.path == '/api/sync':
            data_str = raw.decode('utf-8')
            try:
                incoming = json.loads(data_str)
                if _PG_AVAIL and DATABASE_URL:
                    existing = json.loads(_db_get() or '{}')
                else:
                    try:
                        with open(SYNC_PATH, 'r', encoding='utf-8') as f2:
                            existing = json.loads(f2.read())
                    except Exception:
                        existing = {}
                # -- 空 body 保护：整体为空或只有极少字段的写入直接拒绝 --
                # 防止未登录/初始化态的浏览器把空 localStorage 推上来清空全库。
                if isinstance(incoming, dict):
                    _incoming_keys = set(k for k,v in incoming.items() if v not in (None, {}, [], ''))
                    # 保底核心字段：只要 existing 有 PROJECTS/USER_PROFILES/OPS_ENT 之一，
                    # 而 incoming 三个都为空，就判定为空写入并拒绝。
                    _core_had = any(existing.get(k) for k in ('PROJECTS','USER_PROFILES','OPS_ENT'))
                    _core_incoming = any(incoming.get(k) for k in ('PROJECTS','USER_PROFILES','OPS_ENT'))
                    if _core_had and not _core_incoming:
                        print('[sync] rejected empty write (keys=%r)' % (sorted(_incoming_keys),))
                        resp = json.dumps({'ok': False, 'rejected': 'empty-payload'}).encode()
                        self.send_response(200)
                        self.send_header('Content-Type', 'application/json')
                        self.send_header('Content-Length', str(len(resp)))
                        self.cors(); self.end_headers(); self.wfile.write(resp)
                        return
                # -- RESET generation barrier --
                # admin-reset writes a new RESET_GEN. After that only clients carrying the
                # same RESET_GEN (sessions reloaded after the reset) may write; stale sessions
                # carrying an old/absent gen are rejected wholesale, so they cannot merge the
                # cleared Suizhou runtime data (aiTopics/customTopics/UPLOADS/KB_CHAT) back.
                _srv_gen = existing.get('RESET_GEN')
                if _srv_gen:
                    _cli_gen = incoming.get('RESET_GEN')
                    if _cli_gen != _srv_gen:
                        print('[sync] rejected stale write (client gen=%r != server gen=%r)' % (_cli_gen, _srv_gen))
                        resp = json.dumps({'ok': False, 'rejected': 'stale-generation', 'gen': _srv_gen}).encode()
                        self.send_response(200)
                        self.send_header('Content-Type', 'application/json')
                        self.send_header('Content-Length', str(len(resp)))
                        self.cors(); self.end_headers(); self.wfile.write(resp)
                        return
                # 合并：空列表/空字典不覆盖已有非空数据
                # ── 删除墓碑：政府端删除的项目 key 永久移除。_merge_map 只增不减，
                #    没有墓碑时"删除"永远会被合并复活（删了又回来）。──
                _tomb = set(existing.get('DELETED_PROJECTS') or []) | set(incoming.get('DELETED_PROJECTS') or [])
                _protected = ['OPS_ENT','DEMANDS','KB_CHAT','PENDING_CONFIRMS','KB_CONFIRMS','REPORT_REQUESTS','CITY_ACCOUNTS']
                # REPORT_REQUESTS 按 id 合并且状态只进不退（pending<running<done/failed）
                # 防止管理端旧快照 persist 把流水线已推进的状态倒改回 pending
                _rr_rank = {'pending': 0, 'running': 1, 'failed': 2, 'done': 3}
                def _merge_rr(old_list, new_list):
                    by_id = {r.get('id'): dict(r) for r in (old_list or []) if isinstance(r, dict)}
                    for r in (new_list or []):
                        if not isinstance(r, dict):
                            continue
                        rid = r.get('id')
                        ex = by_id.get(rid)
                        if not ex:
                            by_id[rid] = r
                        elif _rr_rank.get(r.get('status'), 0) >= _rr_rank.get(ex.get('status'), 0):
                            ex.update(r)
                        # 否则丢弃倒退的状态更新，保留服务端已推进的记录
                    return sorted(by_id.values(), key=lambda x: x.get('ts', 0))
                for k, v in incoming.items():
                    if k in _protected and not v and existing.get(k):
                        continue
                    if k == 'REPORT_REQUESTS':
                        existing[k] = _merge_rr(existing.get(k), v)
                        continue
                    if k in ('PROJECTS', 'REPORTSTATE', 'CITY_ACCOUNTS'):
                        # 逐 key/逐字段合并，空值不覆盖非空——防止旧快照把已同步的 RAG/账号冲掉
                        existing[k] = _merge_map(existing.get(k), v)
                        continue
                    existing[k] = v
                # 应用删除墓碑：从 PROJECTS/REPORTSTATE 移除已删项目，并持久化墓碑列表
                if _tomb:
                    existing['DELETED_PROJECTS'] = sorted(_tomb)
                    for _dk in _tomb:
                        for _sect in ('PROJECTS', 'REPORTSTATE', 'PENDING_CONFIRMS', 'UPLOADS', 'KB_FILE_CHUNKS'):
                            if isinstance(existing.get(_sect), dict):
                                existing[_sect].pop(_dk, None)
                data_str = json.dumps(existing, ensure_ascii=False)
            except Exception as e:
                print(f"[sync] merge error: {e}")
            if _PG_AVAIL and DATABASE_URL:
                ok = _db_set(data_str)
                resp = json.dumps({'ok': ok}).encode()
            else:
                try:
                    with open(SYNC_PATH, 'w', encoding='utf-8') as f:
                        f.write(data_str)
                    resp = json.dumps({'ok': True}).encode()
                except Exception as e:
                    resp = json.dumps({'ok': False, 'error': str(e)}).encode()
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Content-Length', str(len(resp)))
            self.cors(); self.end_headers(); self.wfile.write(resp)
            return
        # ── 接口：管理员全量覆写（绕过merge保护，用于重置数据） ──
        if self.path == '/api/admin-reset':
            try:
                body = json.loads(raw)
                secret = body.get('secret', '')
                if secret != 'hxz-reset-2026':
                    resp = json.dumps({'ok': False, 'error': 'unauthorized'}).encode()
                else:
                    new_data = body.get('data')
                    if not isinstance(new_data, dict):
                        resp = json.dumps({'ok': False, 'error': 'data must be dict'}).encode()
                    else:
                        # 每次 reset 生成新的代际标记：此后旧会话(不带此 gen)的 /api/sync 写入被拒绝，
                        # 从根上杜绝已清空的随州数据被旧浏览器快照合并回来。
                        import time as _t
                        _gen = body.get('gen') or ('r' + str(int(_t.time() * 1000)))
                        new_data['RESET_GEN'] = _gen
                        data_str = json.dumps(new_data, ensure_ascii=False)
                        if _PG_AVAIL and DATABASE_URL:
                            ok = _db_set(data_str)
                        else:
                            with open(SYNC_PATH, 'w', encoding='utf-8') as fw:
                                fw.write(data_str)
                            ok = True
                        resp = json.dumps({'ok': ok, 'gen': _gen}).encode()
            except Exception as e:
                resp = json.dumps({'ok': False, 'error': str(e)}).encode()
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(resp)
            return
        # ── 接口0：管理员清理城市智库（强制覆盖kb，绕过材料计数保护） ──
        if self.path == '/api/kb-clean':
            try:
                body = json.loads(raw)
                pkey = body.get('projectKey', '')
                new_kb = body.get('kb')
                if _PG_AVAIL and DATABASE_URL:
                    store = json.loads(_db_get() or '{}')
                else:
                    with open(SYNC_PATH, 'r', encoding='utf-8') as f2:
                        store = json.loads(f2.read())
                projects = store.get('PROJECTS') or {}
                if pkey not in projects:
                    resp = json.dumps({'ok': False, 'error': 'project not found'}).encode()
                elif not isinstance(new_kb, list):
                    resp = json.dumps({'ok': False, 'error': 'kb must be a list'}).encode()
                else:
                    projects[pkey]['kb'] = new_kb
                    data_str = json.dumps(store, ensure_ascii=False)
                    if _PG_AVAIL and DATABASE_URL:
                        _db_set(data_str)
                    else:
                        with open(SYNC_PATH, 'w', encoding='utf-8') as fw:
                            fw.write(data_str)
                    total = sum(len(t.get('known', [])) for t in new_kb)
                    resp = json.dumps({'ok': True, 'total': total}).encode()
            except Exception as e:
                resp = json.dumps({'ok': False, 'error': str(e)}).encode()
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(resp)
            return
        # ── 接口1：管理员上传文档，补充/修正城市智库 ──
        if self.path == '/api/kb-upload':
            try:
                body = json.loads(raw)
                city = body.get('city', '')
                pkey = body.get('projectKey', '')
                topic = body.get('topic', '')          # 目标主题名
                text = body.get('text', '')            # 文档全文（前端已提取或纯文本）
                filename = body.get('filename', '上传文档')
                mode = body.get('mode', 'append')       # append=补充 / replace=修正覆盖该主题
                # 后端文档解析：传 fileB64 时服务端提取文本（PDF/Word），优先于 text
                file_b64 = body.get('fileB64', '')
                if file_b64:
                    try:
                        _fb = base64.b64decode(file_b64)
                    except Exception as _e:
                        resp = json.dumps({'ok': False, 'error': '文件解码失败：' + str(_e)}).encode()
                        self.send_response(200); self.send_header('Content-Type', 'application/json')
                        self.send_header('Content-Length', str(len(resp)))
                        self.cors(); self.end_headers(); self.wfile.write(resp); return
                    _ext_text, _ext_err = _extract_doc_text(filename, _fb)
                    if _ext_err:
                        resp = json.dumps({'ok': False, 'error': _ext_err}).encode()
                        self.send_response(200); self.send_header('Content-Type', 'application/json')
                        self.send_header('Content-Length', str(len(resp)))
                        self.cors(); self.end_headers(); self.wfile.write(resp); return
                    text = _ext_text
                if _PG_AVAIL and DATABASE_URL:
                    store = json.loads(_db_get() or '{}')
                else:
                    with open(SYNC_PATH, 'r', encoding='utf-8') as f2:
                        store = json.loads(f2.read())
                projects = store.get('PROJECTS') or {}
                # 定位项目：优先 projectKey，其次按 city 匹配
                if pkey and pkey in projects:
                    key = pkey
                else:
                    key = next((k for k, v in projects.items()
                                if isinstance(v, dict) and v.get('city') == city), None)
                if not key:
                    resp = json.dumps({'ok': False, 'error': 'project not found'}).encode()
                else:
                    p = projects[key]
                    if not isinstance(p.get('kb'), list):
                        p['kb'] = []
                    # origin: admin(管理员) / user(前端用户)；nature: support(佐证) / fix(修正/修改) / confirm(确认)
                    origin = body.get('origin', 'admin')
                    nature = body.get('nature', 'support')   # 默认佐证
                    import time as _tt
                    _ts = int(_tt.time() * 1000)
                    # 切分文档为结构化 chunks（每条带来源标记，前端据此着色）
                    import re as _re
                    parts = [x.strip() for x in _re.split(r'\n\s*\n', text) if len(x.strip()) >= 20]
                    chunks = []
                    for para in parts:
                        para = _re.sub(r'\s+', ' ', para)[:420]
                        if _is_noise_chunk(para):   # 过滤附注/免责/页码等无价值模板话
                            continue
                        chunks.append({'text': para, 'origin': origin, 'nature': nature,
                                       'src': filename, 'ts': _ts})
                    # 找/建目标主题
                    tp = next((t for t in p['kb'] if t.get('t') == topic), None)
                    if not tp:
                        tp = {'icon': '📎', 't': topic or '管理员补充资料',
                              'sub': '', 'tag': '', 'known': [], 'calls': []}
                        p['kb'].append(tp)
                    # 就地匹配：新 chunk 与已有条目相似度高 → 在已有条目上打标注，不新增
                    def _strip_tag(s):
                        # 剥掉结构化导入的 [对象ID|状态] 前缀，避免稀释相似度
                        return _re.sub(r'^\s*[\[【][^\]】]{0,50}[\]】]\s*', '', str(s or ''))
                    def _norm(s):
                        return _re.sub(r'[\s，,。、；;：:（）()【】\[\]|—\-]+', '', _strip_tag(s))
                    def _chunk_text(k):
                        return k.get('text', '') if isinstance(k, dict) else str(k)
                    def _sal_nums(s):
                        # 显著数字：带小数点 / 百分数 / ≥3位非年份——用于"同一指标同一数值"佐证通道
                        out = set()
                        for m in _re.findall(r'\d+(?:\.\d+)?%?', _strip_tag(s)):
                            if m.endswith('%') or '.' in m:
                                out.add(m.rstrip('%') + ('%' if m.endswith('%') else ''))
                            elif len(m) >= 3 and not (1900 <= int(m) <= 2100):
                                out.add(m)
                        return out
                    def _similar(a, b):
                        a, b = _norm(a), _norm(b)
                        if not a or not b:
                            return 0.0
                        # 双向包含 or 字符级重叠比例
                        if a in b or b in a:
                            return 0.9
                        short, long = (a, b) if len(a) <= len(b) else (b, a)
                        # 滑窗取 short 的若干 8-gram 命中率
                        grams = [short[x:x+8] for x in range(0, max(1, len(short) - 7), 4)] or [short]
                        hit = sum(1 for g in grams if g and g in long)
                        return hit / len(grams)
                    _IND_WORDS = ('生产总值', 'GDP', '增加值', '产值', '产量', '增速', '增长',
                                  '配套率', '城镇化率', '常住人口', '进出口', '进口', '出口',
                                  '财政收入', '投资', '用电', '客运', '货运', '面积', '营收',
                                  '利润', '税收', '就业', '工资', '房价', '地价')
                    def _ind_hit(a, b):
                        # 指标核心词共现：两段文本含同一个指标词才允许数字佐证通道生效
                        return any(w in a and w in b for w in _IND_WORDS)
                    def _is_strong_num(n):
                        # 高精度数字（1502.81 / 75287）撞上即近乎必为同一事实；
                        # 低精度（2.6 / 10% / 100）必须再过指标词门禁
                        d = n.rstrip('%').replace('.', '')
                        return ('.' in n and len(d) >= 4) or ('.' not in n and not n.endswith('%') and len(d) >= 5)

                    matched_ct = 0
                    if mode == 'replace':
                        base = [k for k in (tp.get('known') or [])
                                if isinstance(k, str) or (isinstance(k, dict) and k.get('origin') == 'ai')]
                        tp['known'] = base + chunks
                    else:
                        # 跨主题匹配：原文（如GDP在园区主题、产业事实在产业链主题）分布在
                        # 不同主题，佐证应挂到全库任何一条命中原文上；未命中才追加到目标主题
                        norm_known = []
                        for _tpc in p['kb']:
                            _kn = _tpc.get('known') or []
                            for _i, _k in enumerate(_kn):
                                if not isinstance(_k, dict):
                                    _kn[_i] = {'text': _k, 'origin': 'ai', 'nature': 'base'}
                            _tpc['known'] = _kn
                            norm_known.extend(_kn)
                        leftover = []
                        for nc in chunks:
                            best, best_score = None, 0.55   # 相似度阈值
                            nc_nums = _sal_nums(nc['text'])
                            for ek in norm_known:
                                et = ek.get('text', '')
                                sc = _similar(nc['text'], et)
                                # 数字佐证通道：同一显著数值 → 视为佐证同一事实
                                if sc <= best_score and nc_nums:
                                    # 数字佐证：高精度数字撞上直接挂；低精度数字须指标词共现，
                                    # 防止碰巧同数值的不同指标误挂（增速10% vs 占有率10%、2.6亿 vs 2.6万亿）
                                    shared = nc_nums & _sal_nums(et)
                                    if shared:
                                        if any(_is_strong_num(n) for n in shared):
                                            sc = max(sc, 0.8)   # 高精度数值（如1502.81、75287）
                                        elif _ind_hit(nc['text'], et):
                                            sc = max(sc, 0.7)   # 低精度数值+指标词命中
                                if sc > best_score:
                                    best, best_score = ek, sc
                            if best is not None:
                                # 就地标注到已有条目
                                best.setdefault('annotations', []).append({
                                    'origin': origin, 'nature': nature, 'src': filename,
                                    'text': nc['text'], 'ts': _ts})
                                matched_ct += 1
                            else:
                                leftover.append(nc)
                        tp['known'] = (tp.get('known') or []) + leftover
                    # 记录上传审计
                    p.setdefault('kbUploads', []).append({
                        'filename': filename, 'topic': topic, 'mode': mode,
                        'chunks': len(chunks), 'ts': int(__import__('time').time() * 1000),
                        'by': body.get('by', '管理员'),
                        'source': body.get('source', 'file')})
                    store['PROJECTS'] = projects
                    out_str = json.dumps(store, ensure_ascii=False)
                    if _PG_AVAIL and DATABASE_URL:
                        _db_set(out_str)
                    else:
                        with open(SYNC_PATH, 'w', encoding='utf-8') as f:
                            f.write(out_str)
                    resp = json.dumps({'ok': True, 'chunks': len(chunks),
                                       'matched': matched_ct, 'chars': len(text),
                                       'topic': tp['t'], 'total': len(tp['known'])}).encode()
            except Exception as e:
                resp = json.dumps({'ok': False, 'error': str(e)}).encode()
            self.send_response(200); self.send_header('Content-Type', 'application/json')
            self.send_header('Content-Length', str(len(resp)))
            self.cors(); self.end_headers(); self.wfile.write(resp); return

        # ── 接口2：城市智库版本管理（多次报告按时间快照+覆盖）──
        if self.path == '/api/kb-version':
            try:
                body = json.loads(raw)
                action = body.get('action', 'snapshot')  # snapshot=存快照 / rollback=回滚到某版本
                city = body.get('city', ''); pkey = body.get('projectKey', '')
                if _PG_AVAIL and DATABASE_URL:
                    store = json.loads(_db_get() or '{}')
                else:
                    with open(SYNC_PATH, 'r', encoding='utf-8') as f2:
                        store = json.loads(f2.read())
                projects = store.get('PROJECTS') or {}
                key = pkey if (pkey and pkey in projects) else next(
                    (k for k, v in projects.items() if isinstance(v, dict) and v.get('city') == city), None)
                if not key:
                    resp = json.dumps({'ok': False, 'error': 'project not found'}).encode()
                else:
                    import time as _t
                    p = projects[key]
                    p.setdefault('kbVersions', [])
                    if action == 'snapshot':
                        # 保存当前 kb 为一个带时间戳的历史版本
                        tot = sum(len(t.get('known', [])) for t in (p.get('kb') or []))
                        p['kbVersions'].append({
                            'ver': len(p['kbVersions']) + 1,
                            'ts': int(_t.time() * 1000),
                            'label': body.get('label', ''),
                            'chunks': tot,
                            'kb': json.loads(json.dumps(p.get('kb') or []))  # 深拷贝
                        })
                        # 只保留最近 10 个版本
                        p['kbVersions'] = p['kbVersions'][-10:]
                        resp = json.dumps({'ok': True, 'ver': p['kbVersions'][-1]['ver'],
                                           'versions': len(p['kbVersions'])}).encode()
                    elif action == 'rollback':
                        ver = body.get('ver')
                        target = next((v for v in p['kbVersions'] if v.get('ver') == ver), None)
                        if target:
                            p['kb'] = json.loads(json.dumps(target['kb']))
                            p['kbInitTs'] = int(_t.time() * 1000)
                            resp = json.dumps({'ok': True, 'restored': ver}).encode()
                        else:
                            resp = json.dumps({'ok': False, 'error': 'version not found'}).encode()
                    else:
                        resp = json.dumps({'ok': False, 'error': 'unknown action'}).encode()
                    if b'"ok": true' in resp or b'"ok":true' in resp:
                        store['PROJECTS'] = projects
                        out_str = json.dumps(store, ensure_ascii=False)
                        if _PG_AVAIL and DATABASE_URL:
                            _db_set(out_str)
                        else:
                            with open(SYNC_PATH, 'w', encoding='utf-8') as f:
                                f.write(out_str)
            except Exception as e:
                resp = json.dumps({'ok': False, 'error': str(e)}).encode()
            self.send_response(200); self.send_header('Content-Type', 'application/json')
            self.send_header('Content-Length', str(len(resp)))
            self.cors(); self.end_headers(); self.wfile.write(resp); return

        # ── 接口3：管理端「推送到RAG」按钮 → 给已完成申请打 pushRequested 标记 ──
        # 本地议程轮询器消费该标记，执行 sync_to_kb.py 完成 RAG 推送 + 城市账号连接。
        if self.path == '/api/report-push-request':
            try:
                body = json.loads(raw)
                city = body.get('city', '')
                if _PG_AVAIL and DATABASE_URL:
                    store = json.loads(_db_get() or '{}')
                else:
                    with open(SYNC_PATH, 'r', encoding='utf-8') as f2:
                        store = json.loads(f2.read())
                reqs = store.get('REPORT_REQUESTS') or []
                done = [r for r in reqs if isinstance(r, dict)
                        and r.get('city') == city and r.get('status') == 'done']
                target = (done[-1] if done else
                          (reqs[-1] if reqs else None))
                rid = None
                if target and isinstance(target, dict):
                    target['pushRequested'] = True
                    target['pushRequestedTs'] = int(time.time() * 1000)
                    target['pushed'] = False
                    rid = target.get('id')
                    store['REPORT_REQUESTS'] = reqs
                    out_str = json.dumps(store, ensure_ascii=False)
                    if _PG_AVAIL and DATABASE_URL:
                        _db_set(out_str)
                    else:
                        with open(SYNC_PATH, 'w', encoding='utf-8') as f:
                            f.write(out_str)
                resp = json.dumps({'ok': bool(rid), 'id': rid, 'city': city}).encode()
            except Exception as e:
                resp = json.dumps({'ok': False, 'error': str(e)}).encode()
            self.send_response(200); self.send_header('Content-Type', 'application/json')
            self.send_header('Content-Length', str(len(resp)))
            self.cors(); self.end_headers(); self.wfile.write(resp); return

        if self.path != '/api/kb-chat':
            self.send_error(404); return

        body   = json.loads(raw)
        q      = body.get("question", "").strip()
        chunks = body.get("chunks", [])
        city   = body.get("city", "")
        stream = body.get("stream", True)
        mode   = body.get("mode", "full")
        history = body.get("history", [])  # 多轮上下文：[{role:'user'|'assistant', content:'...'}]
        # 用户初始化(onboarding)招商偏好：后续所有分析的最高优先级约束
        # prefs = {"options": "选项类偏好文本", "custom": "用户自定义输入文本"}
        prefs  = body.get("prefs", {}) or {}

        if mode not in ("chat", "draft", "full", "suggest", "research", "chain", "topics", "funnel", "onboard_options", "city_profile"):
            mode = "full"

        if not q:
            self.send_error(400, "question required"); return

        # 拼 RAG context —— 真实检索：按问题与语料词重合度打分取 TopK，而非盲切前 N
        # 报告类 mode（full/research/chain/funnel）吃进更多语料，对话类少吃
        # city_profile 要横跨全量语料刻画城市，吃进的片段数最多
        _top_k = 40 if mode == "city_profile" else (12 if mode in ("full", "research", "chain", "funnel") else 6)
        if mode == "city_profile":
            # 画像要同时覆盖 经济/结构/产业/园区/链主/成本/偏好 七个面，
            # 单条问题("城市画像…")与 GDP、三次产业结构等语料几乎无词重合 →
            # 单查询检索会整组漏掉经济数据。故按面分别检索再合并去重，保证每个面都有料。
            _facets = [
                "GDP 地区生产总值 总量 增速 人均 财政收入 经济体量",
                "三次产业结构 占比 规上工业增加值 工业占比 第二产业 第三产业",
                "主导产业 产业链 产值 规模 全国占有率 集群 龙头产业",
                "产业园区 经开区 高新区 承载 标准厂房 用地 亩数 入驻率 土地",
                "链主企业 龙头企业 产能 营收 在建项目 配套企业 名录",
                "配套率 成本 人力 物流 交通 competitor 竞争城市 对比",
                "招商方向 优先招引 目标企业 投资强度 政策 补贴 专项资金 领导关注",
                "缺口 薄弱 缺失 短板 数据边界 待补",
            ]
            _per = max(6, _top_k // len(_facets) + 3)
            _seen, used, _scored = set(), [], []
            for _fq in _facets:
                _sel, _sc = _retrieve_chunks(_fq, chunks, top_k=_per, min_score=1)
                for _s, _c in _sc:
                    _cid = id(_c) if not isinstance(_c, dict) else (_c.get("id") or id(_c))
                    if _cid in _seen:
                        continue
                    _seen.add(_cid); used.append(_c); _scored.append((_s, _c))
            # 兜底：分面全未命中时退回原单查询逻辑，避免画像无料可吃
            if not used:
                used, _scored = _retrieve_chunks(q, chunks, top_k=_top_k, min_score=1)
        else:
            used, _scored = _retrieve_chunks(q, chunks, top_k=_top_k, min_score=1)
        # 注意：本方法体内(第729行附近)另有一个局部 def _chunk_text，会遮蔽模块级同名函数，
        # 导致此处 genexpr 引用报 NameError。故内联提取文本，不调用 _chunk_text。
        def _ctext(c):
            if isinstance(c, str): return c
            if isinstance(c, dict): return str(c.get("text", "") or "")
            return ""
        ctx = "\n\n".join(
            f"[{(c.get('cite','') if isinstance(c, dict) else '')}·{(c.get('topic','') if isinstance(c, dict) else '')}]\n{_ctext(c)}"
            for c in used
        ) or "暂无检索到相关内容，请基于通用招商知识回答。"

        # —— RAG 审计日志：证明本次分析实际吃进了哪些数据、其中多少来自用户上传 ——
        # total_available = 前端送来的候选总数；total_chunks = 真实命中并吃进的数量（可能 0~top_k）
        upload_used = [c for c in used if _is_upload_chunk(c)]
        audit_log({
            "city": city,
            "mode": mode,
            "question": q,
            "total_available": len(chunks),
            "total_chunks": len(used),
            "upload_chunks": len(upload_used),
            "kb_chunks": len(used) - len(upload_used),
            "all_cites": [(c.get("cite", "") if isinstance(c, dict) else "") for c in used],
            "top_scores": [s for s, _c in _scored],
            "upload_samples": [
                {"cite": (c.get("cite", "") if isinstance(c, dict) else ""),
                 "id": (c.get("id", "") if isinstance(c, dict) else ""),
                 "text": _ctext(c)[:120]}
                for c in upload_used
            ],
        })

        if mode == "suggest":
            system_prompt = SYSTEM_SUGGEST
            max_tokens    = 200
        elif mode == "chat":
            system_prompt = SYSTEM_CHAT
            max_tokens    = MAX_TOKENS_CHAT
        elif mode == "draft":
            system_prompt = SYSTEM_DRAFT
            max_tokens    = MAX_TOKENS_DRAFT
        elif mode == "research":
            system_prompt = SYSTEM_RESEARCH
            max_tokens    = MAX_TOKENS_RESEARCH
        elif mode == "chain":
            system_prompt = SYSTEM_RESEARCH_CHAIN
            max_tokens    = MAX_TOKENS_CHAIN
        elif mode == "topics":
            system_prompt = SYSTEM_TOPICS
            max_tokens    = 1200
        elif mode == "funnel":
            system_prompt = SYSTEM_FUNNEL
            max_tokens    = 9500    # 约40家企业结构化 JSON（含三分项评分卡+扩张/派系标注，字段增多）
        elif mode == "onboard_options":
            system_prompt = SYSTEM_ONBOARD
            max_tokens    = 900
        elif mode == "city_profile":
            system_prompt = SYSTEM_CITY_PROFILE
            max_tokens    = 8000    # 结构化画像 JSON：objective 六组 + preference 五维 + 优劣势/缺口
        else:
            system_prompt = SYSTEM_FULL
            max_tokens    = MAX_TOKENS_FULL

        # 多轮上下文：注入最近几轮对话（仅保留合法 role + 非空文本，限制条数防止 context 膨胀）
        hist_msgs = []
        for h in (history or [])[-8:]:
            role = h.get("role", "")
            content = (h.get("content", "") or "").strip()
            if role in ("user", "assistant") and content:
                hist_msgs.append({"role": role, "content": content[:1500]})

        # 用户初始化招商偏好块：最高优先级注入到 user message 顶部（onboard_options 自身在生成偏好，不注入）
        prefs_block = ""
        if mode == "city_profile":
            # 画像模式：问卷是"待刻画的证据/校准项"，不是要服从的约束。
            # 故单独组织成 stated 证据块；缺失时明确告知模型走纯反推 + 标 gaps。
            _opt = (prefs.get("options") or "").strip()
            _cus = (prefs.get("custom") or "").strip()
            if _opt or _cus:
                parts = ["【招商偏好问卷 answers（stated 证据，用于校准反推结论）】"]
                if _opt:
                    parts.append("问卷选择：\n" + _opt[:1500])
                if _cus:
                    parts.append("干部自定义补充：\n" + _cus[:1500])
                parts.append("说明：以上为该城市干部明确作答的偏好，对应维度 source 标 stated；与材料反推冲突时以问卷为准并记入 conflicts。")
                prefs_block = "\n\n".join(parts) + "\n\n"
            else:
                prefs_block = ("【招商偏好问卷：该城市尚未填写】\n"
                               "所有 preference 维度只能从材料反推，source 一律标 inferred，"
                               "并在 gaps 中提示需补填招商偏好问卷。\n\n")
        elif mode != "onboard_options":
            _opt = (prefs.get("options") or "").strip()
            _cus = (prefs.get("custom") or "").strip()
            if _opt or _cus:
                parts = ["【用户招商偏好——最高优先级，以下分析必须优先满足并贯穿始终】"]
                if _opt:
                    parts.append("招商偏好选择：\n" + _opt[:1200])
                if _cus:
                    # 用户自定义输入优先级最高，单独强调
                    parts.append("⭐用户自定义强调（优先级最高，须重点体现、不得忽略或弱化）：\n" + _cus[:1500])
                parts.append("要求：产业方向、目标企业、园区匹配、招引策略等所有结论都要与上述偏好一致；如知识片段与用户偏好冲突，以用户偏好为准并说明。")
                prefs_block = "\n\n".join(parts) + "\n\n"

        user_msg = f"{prefs_block}城市：{city}\n\n知识片段：\n{ctx}\n\n问题：{q}"

        payload_dict = {
            "model": MODEL,
            "max_tokens": max_tokens,
            "stream": stream,
            "temperature": 0,
            "messages": [
                {"role": "system", "content": system_prompt},
                *hist_msgs,
                {"role": "user",   "content": user_msg}
            ]
        }
        # v4-pro 默认开启 thinking，CoT 会先吐 reasoning_content 并吃光 max_tokens，
        # 导致 content 迟迟不出/为空。前端只渲染 delta.content、不显示思考链，
        # 所以对所有 mode（含 chat/suggest）一律关闭 thinking，避免问答“出不来答案”。
        payload_dict["thinking"] = {"type": "disabled"}

        # chat 模式改走 Responses API + v4-flash 原生联网搜索（web_search 仅 flash 支持）。
        # 后端把 Responses 流式事件翻译回 chat-completions SSE 格式，前端零改动。
        if mode in ("chat", "onboard_options"):
            resp_payload = {
                "model": MODEL_CHAT,
                "max_output_tokens": max_tokens + 600,  # 预留搜索/工具调用的输出开销
                "stream": stream,
                "instructions": system_prompt,
                "input": [*hist_msgs,
                          {"role": "user", "content": user_msg}],
                "tools": [{"type": "web_search"}],
                "thinking": {"type": "disabled"},
            }
            payload = json.dumps(resp_payload).encode()
            ds_url = DS_RESP_URL
        else:
            payload = json.dumps(payload_dict).encode()
            ds_url = DS_URL

        req = urllib.request.Request(
            ds_url, data=payload, method="POST",
            headers={"Content-Type": "application/json",
                     "Authorization": f"Bearer {DS_KEY}"}
        )

        try:
            _timeout = 150 if mode in ("funnel", "chain", "research") else 90
            with DS_OPENER.open(req, timeout=_timeout) as resp:
                if stream:
                    self.send_response(200)
                    self.send_header("Content-Type", "text/event-stream")
                    self.send_header("Cache-Control", "no-cache")
                    self.cors(); self.end_headers()
                    if mode in ("chat", "onboard_options"):
                        # Responses SSE → chat-completions SSE 翻译层
                        # 修复"先卡住再涌出"：web_search/思考阶段上游长时间不吐 output_text，
                        # 前端旧逻辑只能干等。现把中间状态事件翻译成 reasoning_content 发给前端，
                        # 前端已有 reasoning 分支会显示"深度推理中"进度，同时也起到保活作用。
                        for raw_line in resp:
                            line = raw_line.decode("utf-8", errors="replace").strip()
                            if not line.startswith("data:"): continue
                            data = line[5:].strip()
                            if not data: continue
                            try: ev = json.loads(data)
                            except Exception: continue
                            et = ev.get("type", "")
                            if et == "response.output_text.delta":
                                out = json.dumps({"choices": [{"delta": {"content": ev.get("delta", "")}}]}, ensure_ascii=False)
                                self.wfile.write(f"data: {out}\n\n".encode()); self.wfile.flush()
                            elif et in ("response.web_search_call.in_progress", "response.web_search_call.searching"):
                                out = json.dumps({"choices": [{"delta": {"reasoning_content": "[searching]"}}]}, ensure_ascii=False)
                                self.wfile.write(f"data: {out}\n\n".encode()); self.wfile.flush()
                            elif et == "response.reasoning_text.delta":
                                out = json.dumps({"choices": [{"delta": {"reasoning_content": ev.get("delta", "")}}]}, ensure_ascii=False)
                                self.wfile.write(f"data: {out}\n\n".encode()); self.wfile.flush()
                            elif et in ("response.completed", "response.failed", "response.incomplete"):
                                self.wfile.write(b"data: [DONE]\n\n"); self.wfile.flush()
                    else:
                        while True:
                            line = resp.readline()
                            if not line: break
                            self.wfile.write(line); self.wfile.flush()
                        # 显式发送[DONE]确保前端reader.read()收到done=true
                        self.wfile.write(b"data: [DONE]\n\n"); self.wfile.flush()
                else:
                    data = resp.read()
                    if mode in ("chat", "onboard_options"):
                        # Responses JSON → chat-completions JSON 翻译
                        d = json.loads(data)
                        txt = "".join(
                            c.get("text", "")
                            for i in d.get("output", []) if i.get("type") == "message"
                            for c in i.get("content", []) if c.get("type") == "output_text"
                        )
                        data = json.dumps({"choices": [{"message": {"role": "assistant", "content": txt}}]}, ensure_ascii=False).encode()
                    self.send_response(200)
                    self.send_header("Content-Type", "application/json")
                    self.cors(); self.end_headers(); self.wfile.write(data)

        except urllib.error.HTTPError as e:
            err = e.read().decode(errors="replace")
            print(f"[kb] upstream {e.code}: {err[:200]}")
            self.send_response(e.code)
            self.send_header("Content-Type", "application/json")
            self.cors(); self.end_headers()
            self.wfile.write(json.dumps({"error": err[:400]}).encode())
        except Exception as e:
            print(f"[kb] error: {e}")
            self.send_response(500)
            self.send_header("Content-Type", "application/json")
            self.cors(); self.end_headers()
            self.wfile.write(json.dumps({"error": str(e)}).encode())


if __name__ == "__main__":
    if not DS_KEY:
        print("[kb] WARN: DEEPSEEK_API_KEY 未设置 — 页面可访问，但 AI 问答/报告会报错。请在 Railway Variables 里配置 DEEPSEEK_API_KEY。")
    _init_db()
    import threading

    # 云端(Railway)：单端口，绑 0.0.0.0；政府端=/，管理端=/ops（同一端口路由）
    IS_CLOUD = bool(os.environ.get("PORT"))
    HOST = "0.0.0.0" if IS_CLOUD else "127.0.0.1"

    if not IS_CLOUD:
        # 本地开发：额外起 5051 兼容旧用法
        ops_server = ThreadingHTTPServer((HOST, PORT_OPS), Handler)
        threading.Thread(target=ops_server.serve_forever, daemon=True).start()
        print(f"[kb] 管理端(本地) http://localhost:{PORT_OPS}  (ops)")

    print(f"[kb] 服务启动 http://{HOST}:{PORT_GOV}  政府端=/  管理端=/ops")
    gov_server = ThreadingHTTPServer((HOST, PORT_GOV), Handler)
    gov_server.serve_forever()
